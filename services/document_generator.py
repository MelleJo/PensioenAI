from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

class DocumentGenerator:
    def __init__(self):
        self.logo_path_veldhuis = "assets/veldhuis_logo.png"
        self.logo_path_jip = "assets/jip_logo.png"
        self.BLUE_COLOR = "1B4E9C"  # Veldhuis blue
        self.RED_COLOR = "FF0000"   # Red for company name

    def create_report(self, report_type, content, company_name, date, advisor_name):
        doc = Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        # Add logos
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        
        # Left logo cell
        left_cell = table.cell(0, 0)
        left_logo = left_cell.paragraphs[0]
        left_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        try:
            run = left_logo.add_run()
            run.add_picture(self.logo_path_veldhuis, width=Inches(2))
        except:
            run.add_text("VELDHUIS ADVIES")

        # Right logo cell
        right_cell = table.cell(0, 1)
        right_logo = right_cell.paragraphs[0]
        right_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        try:
            run = right_logo.add_run()
            run.add_picture(self.logo_path_jip, width=Inches(1))
        except:
            run.add_text("JIP financieel")

        doc.add_paragraph()  # Spacing

        # Title section with blue background
        title = ("Advies pensioenregeling conform de Wet Toekomst Pensioenen" 
                if report_type == "advice" 
                else "Inventarisatie overstap naar de Wet Toekomst Pensioenen")
        
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        
        # Add blue background
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{self.BLUE_COLOR}"/>')
        title_paragraph._p.get_or_add_pPr().append(shading_elm)
        title_run.font.color.rgb = RGBColor(255, 255, 255)

        # Company name in red
        company_paragraph = doc.add_paragraph()
        company_run = company_paragraph.add_run(company_name)
        company_run.font.color.rgb = RGBColor(255, 0, 0)
        company_run.font.size = Pt(24)
        company_run.font.bold = True
        company_paragraph.space_after = Pt(12)

        # Report type and metadata
        info_paragraph = doc.add_paragraph()
        info_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        info_paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE
        info_run = info_paragraph.add_run(f"{'Adviesrapportage' if report_type == 'advice' else 'Analyserapportage'}\n")
        info_run.font.size = Pt(12)
        info_paragraph.add_run(f"{date}\n").font.size = Pt(12)
        info_paragraph.add_run(f"{advisor_name}\n").font.size = Pt(12)
        info_paragraph.add_run("Veldhuis Advies / JIP Financieel").font.size = Pt(12)
        info_paragraph.space_after = Pt(24)

        # Content
        heading = doc.add_heading('Samenvatting', level=1)
        heading.style.font.size = Pt(14)
        heading.style.font.bold = True
        
        content_paragraph = doc.add_paragraph()
        content_run = content_paragraph.add_run(content)
        content_run.font.size = Pt(11)

        # Footer
        footer = doc.sections[0].footer
        footer_text = footer.paragraphs[0]
        footer_text.text = "Financiële planning | Pensioenen | Hypotheken | Verzekeringen | Risicomanagement"
        footer_text.alignment = WD_ALIGN_PARAGRAPH.LEFT
        footer_text.style.font.size = Pt(8)
        footer_text.style.font.color.rgb = RGBColor(128, 128, 128)

        return doc

    def save_document(self, doc, filename):
        doc.save(filename)